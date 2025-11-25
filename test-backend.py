from fastapi import FastAPI, HTTPException, Query
from fastapi.responses import FileResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, HttpUrl
from youtube_transcript_api import YouTubeTranscriptApi
from youtube_transcript_api._errors import TranscriptsDisabled, NoTranscriptFound
import re
from io import BytesIO
from typing import Literal
import os

# Document generation imports
from docx import Document
from fpdf import FPDF
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

app = FastAPI(title="YouTube Transcript API Service")

# CORS Configuration
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # ใน production ควรระบุ domain ที่อนุญาต
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Models
class TranscriptRequest(BaseModel):
    video_url: str
    languages: list[str] = ["en", "th"]
    format: Literal["txt", "docx", "pdf"] = "txt"
    include_timestamps: bool = True

class TranscriptResponse(BaseModel):
    video_id: str
    language: str
    language_code: str
    is_generated: bool
    snippet_count: int

# Helper Functions
def extract_video_id(url: str) -> str:
    """ดึง video ID จาก YouTube URL"""
    patterns = [
        r'(?:youtube\.com\/watch\?v=|youtu\.be\/|youtube\.com\/embed\/)([^&\n?#]+)',
        r'^([a-zA-Z0-9_-]{11})$'  # รองรับ video ID โดยตรง
    ]
    
    for pattern in patterns:
        match = re.search(pattern, url)
        if match:
            return match.group(1)
    
    raise ValueError("Invalid YouTube URL or Video ID")

def generate_txt(transcript, include_timestamps: bool) -> BytesIO:
    """สร้างไฟล์ TXT"""
    buffer = BytesIO()
    content = ""
    
    for snippet in transcript.snippets:
        if include_timestamps:
            content += f"[{snippet['start']:.2f}s] {snippet['text']}\n"
        else:
            content += f"{snippet['text']}\n"
    
    buffer.write(content.encode('utf-8'))
    buffer.seek(0)
    return buffer

def generate_docx(transcript, video_id: str, include_timestamps: bool) -> BytesIO:
    """สร้างไฟล์ DOCX"""
    doc = Document()
    
    # เพิ่มหัวเรื่อง
    doc.add_heading(f'YouTube Transcript - {video_id}', 0)
    doc.add_paragraph(f'Language: {transcript.language} ({transcript.language_code})')
    doc.add_paragraph(f'Auto-generated: {transcript.is_generated}')
    doc.add_paragraph('')
    
    # เพิ่มเนื้อหา transcript
    for snippet in transcript.snippets:
        if include_timestamps:
            text = f"[{snippet['start']:.2f}s] {snippet['text']}"
        else:
            text = snippet['text']
        doc.add_paragraph(text)
    
    # บันทึกลง buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generate_pdf(transcript, video_id: str, include_timestamps: bool) -> BytesIO:
    """สร้างไฟล์ PDF ด้วย FPDF (รองรับ Unicode)"""
    pdf = FPDF()
    pdf.add_page()
    
    # ตั้งค่า font (ใช้ DejaVu ที่รองรับ Unicode)
    pdf.add_font('DejaVu', '', '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf', uni=True)
    pdf.set_font('DejaVu', '', 12)
    
    # หัวเรื่อง
    pdf.set_font('DejaVu', '', 16)
    pdf.cell(0, 10, f'YouTube Transcript - {video_id}', ln=True, align='C')
    pdf.ln(5)
    
    pdf.set_font('DejaVu', '', 10)
    pdf.cell(0, 5, f'Language: {transcript.language} ({transcript.language_code})', ln=True)
    pdf.cell(0, 5, f'Auto-generated: {transcript.is_generated}', ln=True)
    pdf.ln(5)
    
    # เนื้อหา transcript
    pdf.set_font('DejaVu', '', 11)
    for snippet in transcript.snippets:
        if include_timestamps:
            text = f"[{snippet['start']:.2f}s] {snippet['text']}"
        else:
            text = snippet['text']
        
        # ใช้ multi_cell เพื่อจัดการข้อความยาว
        pdf.multi_cell(0, 6, text)
    
    # บันทึกลง buffer
    buffer = BytesIO()
    pdf.output(buffer)
    buffer.seek(0)
    return buffer

# API Endpoints
@app.get("/")
async def root():
    return {
        "message": "YouTube Transcript API Service",
        "endpoints": {
            "GET /transcript": "Get transcript and download as file",
            "GET /transcript/info": "Get transcript information only",
            "GET /transcript/languages": "List available languages"
        }
    }

@app.get("/transcript")
async def get_transcript(
    video_url: str = Query(..., description="YouTube video URL or ID"),
    languages: str = Query("en,th", description="Comma-separated language codes"),
    format: Literal["txt", "docx", "pdf"] = Query("txt", description="Output format"),
    include_timestamps: bool = Query(True, description="Include timestamps")
):
    """
    ดึง transcript และส่งกลับเป็นไฟล์
    
    Example:
    /transcript?video_url=dQw4w9WgXcQ&languages=en,th&format=pdf&include_timestamps=true
    """
    try:
        # ดึง video ID
        video_id = extract_video_id(video_url)
        
        # ดึง transcript
        ytt_api = YouTubeTranscriptApi()
        lang_list = [lang.strip() for lang in languages.split(',')]
        transcript = ytt_api.fetch(video_id, languages=lang_list)
        
        # สร้างไฟล์ตาม format
        filename = f"transcript_{video_id}.{format}"
        
        if format == "txt":
            buffer = generate_txt(transcript, include_timestamps)
            media_type = "text/plain"
        elif format == "docx":
            buffer = generate_docx(transcript, video_id, include_timestamps)
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif format == "pdf":
            buffer = generate_pdf(transcript, video_id, include_timestamps)
            media_type = "application/pdf"
        
        return StreamingResponse(
            buffer,
            media_type=media_type,
            headers={
                "Content-Disposition": f"attachment; filename={filename}"
            }
        )
        
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except (TranscriptsDisabled, NoTranscriptFound) as e:
        raise HTTPException(status_code=404, detail=f"Transcript not found: {str(e)}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Internal error: {str(e)}")

@app.get("/transcript/info")
async def get_transcript_info(
    video_url: str = Query(..., description="YouTube video URL or ID"),
    languages: str = Query("en,th", description="Comma-separated language codes")
):
    """
    ดึงข้อมูล transcript เท่านั้น (ไม่สร้างไฟล์)
    """
    try:
        video_id = extract_video_id(video_url)
        ytt_api = YouTubeTranscriptApi()
        lang_list = [lang.strip() for lang in languages.split(',')]
        transcript = ytt_api.fetch(video_id, languages=lang_list)
        
        return TranscriptResponse(
            video_id=video_id,
            language=transcript.language,
            language_code=transcript.language_code,
            is_generated=transcript.is_generated,
            snippet_count=len(list(transcript.snippets))
        )
        
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except (TranscriptsDisabled, NoTranscriptFound) as e:
        raise HTTPException(status_code=404, detail=f"Transcript not found: {str(e)}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Internal error: {str(e)}")

@app.get("/transcript/languages")
async def get_available_languages(
    video_url: str = Query(..., description="YouTube video URL or ID")
):
    """
    แสดงรายการภาษาที่มีให้สำหรับ video นี้
    """
    try:
        video_id = extract_video_id(video_url)
        ytt_api = YouTubeTranscriptApi()
        transcript_list = ytt_api.list(video_id)
        
        languages = []
        for transcript in transcript_list:
            languages.append({
                "language": transcript.language,
                "language_code": transcript.language_code,
                "is_generated": transcript.is_generated,
                "is_translatable": transcript.is_translatable
            })
        
        return {
            "video_id": video_id,
            "available_languages": languages
        }
        
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Internal error: {str(e)}")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
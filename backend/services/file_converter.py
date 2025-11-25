"""
File Converter Service - บริการสำหรับแปลง transcript เป็นไฟล์ต่างๆ
รองรับ: TXT, PDF, DOCX
"""

import os
import tempfile
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


class FileConverter:
    """Service สำหรับแปลง transcript เป็นไฟล์ต่างๆ"""
    
    def __init__(self):
        self.temp_dir = tempfile.gettempdir()
    
    def to_txt(self, transcript, include_timestamps: bool = True) -> str:
        """
        แปลง transcript เป็นไฟล์ TXT
        
        Args:
            transcript: FetchedTranscript object
            include_timestamps: รวม timestamps หรือไม่
        
        Returns:
            Path ของไฟล์ที่สร้าง
        """
        temp_file = tempfile.NamedTemporaryFile(
            mode='w',
            suffix='.txt',
            delete=False,
            encoding='utf-8'
        )
        
        try:
            # เขียน header
            temp_file.write(f"YouTube Transcript\n")
            temp_file.write(f"{'=' * 60}\n")
            temp_file.write(f"Video ID: {transcript.video_id}\n")
            temp_file.write(f"Language: {transcript.language} ({transcript.language_code})\n")
            temp_file.write(f"Auto-generated: {'Yes' if transcript.is_generated else 'No'}\n")
            temp_file.write(f"{'=' * 60}\n\n")
            
            # เขียนเนื้อหา
            for snippet in transcript:
                if include_timestamps:
                    start_time = self._format_timestamp(snippet.start)
                    end_time = self._format_timestamp(snippet.start + snippet.duration)
                    temp_file.write(f"[{start_time} - {end_time}]\n")
                temp_file.write(f"{snippet.text}\n\n")
            
            temp_file.close()
            return temp_file.name
        
        except Exception as e:
            temp_file.close()
            if os.path.exists(temp_file.name):
                os.unlink(temp_file.name)
            raise Exception(f"ไม่สามารถสร้างไฟล์ TXT ได้: {str(e)}")
    
    def to_pdf(self, transcript, include_timestamps: bool = True) -> str:
        """
        แปลง transcript เป็นไฟล์ PDF
        
        Args:
            transcript: FetchedTranscript object
            include_timestamps: รวม timestamps หรือไม่
        
        Returns:
            Path ของไฟล์ที่สร้าง
        """
        temp_file = tempfile.NamedTemporaryFile(
            suffix='.pdf',
            delete=False
        )
        temp_file.close()
        
        try:
            doc = SimpleDocTemplate(
                temp_file.name,
                pagesize=A4,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=18
            )
            
            # สร้าง styles
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                textColor='#333333',
                spaceAfter=12
            )
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=12,
                textColor='#666666',
                spaceAfter=6
            )
            body_style = ParagraphStyle(
                'CustomBody',
                parent=styles['Normal'],
                fontSize=10,
                textColor='#000000',
                spaceAfter=12,
                leading=14
            )
            timestamp_style = ParagraphStyle(
                'Timestamp',
                parent=styles['Normal'],
                fontSize=9,
                textColor='#888888',
                spaceAfter=4
            )
            
            story = []
            
            # Header
            story.append(Paragraph("YouTube Transcript", title_style))
            story.append(Spacer(1, 0.2 * inch))
            
            # Metadata
            story.append(Paragraph(f"<b>Video ID:</b> {transcript.video_id}", heading_style))
            story.append(Paragraph(
                f"<b>Language:</b> {transcript.language} ({transcript.language_code})",
                heading_style
            ))
            story.append(Paragraph(
                f"<b>Auto-generated:</b> {'Yes' if transcript.is_generated else 'No'}",
                heading_style
            ))
            story.append(Spacer(1, 0.3 * inch))
            
            # Content
            for snippet in transcript:
                if include_timestamps:
                    start_time = self._format_timestamp(snippet.start)
                    end_time = self._format_timestamp(snippet.start + snippet.duration)
                    story.append(Paragraph(
                        f"[{start_time} - {end_time}]",
                        timestamp_style
                    ))
                
                # Escape HTML characters และแปลงเป็น paragraph
                text = snippet.text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                story.append(Paragraph(text, body_style))
                story.append(Spacer(1, 0.1 * inch))
            
            doc.build(story)
            return temp_file.name
        
        except Exception as e:
            if os.path.exists(temp_file.name):
                os.unlink(temp_file.name)
            raise Exception(f"ไม่สามารถสร้างไฟล์ PDF ได้: {str(e)}")
    
    def to_docx(self, transcript, include_timestamps: bool = True) -> str:
        """
        แปลง transcript เป็นไฟล์ DOCX
        
        Args:
            transcript: FetchedTranscript object
            include_timestamps: รวม timestamps หรือไม่
        
        Returns:
            Path ของไฟล์ที่สร้าง
        """
        temp_file = tempfile.NamedTemporaryFile(
            suffix='.docx',
            delete=False
        )
        temp_file.close()
        
        try:
            doc = Document()
            
            # Header
            title = doc.add_heading('YouTube Transcript', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Metadata
            doc.add_paragraph(f"Video ID: {transcript.video_id}")
            doc.add_paragraph(f"Language: {transcript.language} ({transcript.language_code})")
            doc.add_paragraph(f"Auto-generated: {'Yes' if transcript.is_generated else 'No'}")
            doc.add_paragraph("")  # Empty line
            
            # Content
            for snippet in transcript:
                if include_timestamps:
                    start_time = self._format_timestamp(snippet.start)
                    end_time = self._format_timestamp(snippet.start + snippet.duration)
                    timestamp_para = doc.add_paragraph(f"[{start_time} - {end_time}]")
                    timestamp_para.style.font.size = Pt(9)
                    timestamp_para.style.font.color.rgb = None  # Gray color
                
                text_para = doc.add_paragraph(snippet.text)
                text_para.style.font.size = Pt(11)
                doc.add_paragraph("")  # Empty line between snippets
            
            doc.save(temp_file.name)
            return temp_file.name
        
        except Exception as e:
            if os.path.exists(temp_file.name):
                os.unlink(temp_file.name)
            raise Exception(f"ไม่สามารถสร้างไฟล์ DOCX ได้: {str(e)}")
    
    def _format_timestamp(self, seconds: float) -> str:
        """
        แปลงวินาทีเป็นรูปแบบ timestamp (HH:MM:SS)
        
        Args:
            seconds: จำนวนวินาที
        
        Returns:
            String ในรูปแบบ HH:MM:SS หรือ MM:SS
        """
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        
        if hours > 0:
            return f"{hours:02d}:{minutes:02d}:{secs:02d}"
        else:
            return f"{minutes:02d}:{secs:02d}"

